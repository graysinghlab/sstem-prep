
import os, smtplib, ssl, io
from email.mime.text import MIMEText
from datetime import datetime
import pandas as pd
import streamlit as st
import yaml
import matplotlib.pyplot as plt
from docx import Document

BASE = os.path.dirname(__file__)
TRACKER = os.path.join(BASE, "tracker.csv")
CFG_PATH = os.path.join(BASE, "config.yaml")

st.set_page_config(page_title="S-STEM Preparatory Phase", layout="wide")

@st.cache_data
def load_config():
    import yaml, os
    with open(CFG_PATH, "r") as f:
        return yaml.safe_load(f)

def save_config(cfg):
    import yaml
    with open(CFG_PATH, "w") as f:
        yaml.safe_dump(cfg, f)

@st.cache_data
def load_tracker():
    import pandas as pd
    df = pd.read_csv(TRACKER, dtype=str).fillna("")
    return df

def save_tracker(df):
    df.to_csv(TRACKER, index=False)

def send_email(cfg, subject, html):
    import ssl, smtplib
    from email.mime.text import MIMEText
    smtp = cfg["smtp"]
    recipients = cfg["recipients"]
    if "YOUR_EMAIL_HERE" in recipients or smtp["username"] == "YOUR_SMTP_USERNAME":
        st.warning("Configure recipients and SMTP credentials in the Settings tab before sending alerts.")
        return
    msg = MIMEText(html, "html")
    msg["Subject"] = subject
    msg["From"] = smtp["username"]
    msg["To"] = ",".join(recipients)
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL(smtp["host"], smtp["port"], context=context) as server:
        server.login(smtp["username"], smtp["password"])
        server.sendmail(msg["From"], recipients, msg.as_string())

def donut_png(percent: int):
    import matplotlib.pyplot as plt
    import io
    fig, ax = plt.subplots(figsize=(1.2,1.2))
    ax.pie([percent, 100-percent], startangle=90, counterclock=False,
           wedgeprops=dict(width=0.3, edgecolor='white'),
           colors=[(0.10,0.45,0.90), (0.94,0.96,0.99)])
    ax.text(0,0,f"{percent}%", ha="center", va="center", fontsize=10, color=(0.08,0.09,0.1), fontweight="bold")
    ax.set(aspect="equal")
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=200, bbox_inches="tight", transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf

def compute_stats(df):
    total = len(df)
    done = (df["Status"]=="Complete").sum()
    overall_pct = int(round(100*done/max(total,1)))
    cats = sorted(df["Category"].unique().tolist())
    cat_stats = []
    for cat in cats:
        sub = df[df["Category"]==cat]
        cdone = (sub["Status"]=="Complete").sum()
        pct = int(round(100*cdone/max(len(sub),1)))
        cat_stats.append((cat, pct, cdone, len(sub)))
    return overall_pct, done, total, cat_stats

if "category_filter" not in st.session_state:
    st.session_state["category_filter"] = "(All)"

cfg = load_config()
st.markdown("""
    <style>
    .card { background:#fff; border:1px solid #e9edf3; border-radius:14px; padding:14px; }
    .muted { color:#64748b; font-size:13px; }
    .kpi { font-size:22px; font-weight:700; color:#0f172a; }
    .btn-primary { background:#0b57d0; color:#fff; padding:6px 12px; border-radius:10px; text-decoration:none; }
    .chip { display:inline-block; padding:6px 10px; border-radius:20px; border:1px solid #e5e7eb; margin:0 6px 6px 0; font-size:12px; }
    .chip-active { background:#ebf3ff; border-color:#cfe0ff; color:#0b57d0; }
    </style>
""", unsafe_allow_html=True)

st.title(cfg["app"]["title"])
tab_dash, tab_tasks, tab_settings = st.tabs(["Dashboard","Tasks","Settings"])

with tab_dash:
    df = load_tracker()
    overall_pct, done, total, cat_stats = compute_stats(df)
    c1, c2 = st.columns([5,1])
    with c1:
        st.markdown("#### Overall Progress")
        st.progress(overall_pct/100.0, text=f"{done} of {total} complete · {overall_pct}%")
    with c2:
        st.download_button("Export CSV", data=df.to_csv(index=False).encode("utf-8"), file_name="tracker.csv", use_container_width=True)

    st.markdown("---")
    grid_cols = st.columns(4)
    for i, (cat, pct, cdone, clen) in enumerate(cat_stats):
        with grid_cols[i % 4]:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            a,b = st.columns([1,2])
            with a:
                st.image(donut_png(pct), use_column_width=False)
            with b:
                st.markdown(f"**{cat}**")
                st.caption(f"{cdone} of {clen} complete")
                if st.button("Open tasks", key=f"go_{cat}"):
                    st.session_state["category_filter"] = cat
                    st.experimental_rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Quick View")
    qf = st.session_state["category_filter"]
    st.caption(f"Showing: {qf}")
    qdf = df if qf == "(All)" else df[df["Category"]==qf]
    st.dataframe(qdf[["Category","Reviewer Concern","Owner","Status","Due Date","Notes / Evidence Link"]],
                 use_container_width=True, hide_index=True)

with tab_tasks:
    df = load_tracker()
    fc1, fc2, fc3, fc4 = st.columns([2,2,2,2])
    with fc1:
        cats = ["(All)"] + sorted(df["Category"].unique().tolist())
        cat = st.selectbox("Category", cats, index=cats.index(st.session_state.get("category_filter","(All)")))
    with fc2:
        owners = ["(Any)"] + sorted([x for x in df["Owner"].unique().tolist() if x])
        owner = st.selectbox("Owner", owners, index=0)
    with fc3:
        status_opts = ["(All)","Not Started","In Progress","Complete"]
        stat = st.selectbox("Status", status_opts, index=0)
    with fc4:
        search = st.text_input("Search", placeholder="Find by concern or revision")

    view = df.copy()
    if cat != "(All)":
        view = view[view["Category"]==cat]
    if owner != "(Any)":
        view = view[view["Owner"]==owner]
    if stat != "(All)":
        view = view[view["Status"]==stat]
    if search.strip():
        mask = view["Reviewer Concern"].str.contains(search, case=False) | view["Revision Planned"].str.contains(search, case=False)
        view = view[mask]

    st.markdown("#### Edit Tasks")
    edited = st.data_editor(
        view,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        column_config={
            "Due Date": st.column_config.DateColumn(format="YYYY-MM-DD"),
            "Completion Date": st.column_config.TextColumn(),
            "Notes / Evidence Link": st.column_config.LinkColumn()
        }
    )

    if st.button("Save Changes", type="primary"):
        base = df.merge(
            edited[["Reviewer Concern","Revision Planned","Checklist Ref.","Owner","Status","Due Date","Notes / Evidence Link"]],
            on=["Reviewer Concern","Revision Planned","Checklist Ref."],
            how="left", suffixes=("","_new")
        )
        transitions = []
        for i,row in base.iterrows():
            prev = row["Status"]
            new = row["Status_new"] if pd.notna(row["Status_new"]) else row["Status"]
            if pd.notna(row["Owner_new"]): base.at[i,"Owner"] = row["Owner_new"]
            if pd.notna(row["Status_new"]): base.at[i,"Status"] = row["Status_new"]
            if pd.notna(row["Due Date_new"]): base.at[i,"Due Date"] = row["Due Date_new"]
            if pd.notna(row["Notes / Evidence Link_new"]): base.at[i,"Notes / Evidence Link"] = row["Notes / Evidence Link_new"]
            base.at[i,"Last Updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            if prev != "Complete" and new == "Complete":
                base.at[i,"Completion Date"] = datetime.now().strftime("%Y-%m-%d")
                transitions.append(i)
            elif prev == "Complete" and new != "Complete":
                base.at[i,"Completion Date"] = ""

        for c in [c for c in base.columns if c.endswith("_new")]:
            base.drop(columns=[c], inplace=True)

        base.to_csv(TRACKER, index=False)
        st.success("Saved.")

        cfg = load_config()
        if transitions:
            for _, r in base.loc[transitions].iterrows():
                html = f"""
                <p><b>Task Completed</b></p>
                <p><b>Category:</b> {r['Category']}</p>
                <p><b>Reviewer Concern:</b> {r['Reviewer Concern']}</p>
                <p><b>Revision Planned:</b> {r['Revision Planned']}</p>
                <p><b>Owner:</b> {r['Owner']}</p>
                <p><b>Evidence:</b> {r['Notes / Evidence Link'] or '(none)'} </p>
                """
                try:
                    send_email(cfg, f"S-STEM Prep: Task Complete — {r['Reviewer Concern']}", html)
                except Exception as e:
                    st.warning(f"Email failed (configure SMTP in Settings): {e}")

        cat_done = base.groupby("Category")["Status"].apply(lambda s: (s=='Complete').sum()==len(s))
        if "cat_done_state" not in st.session_state:
            st.session_state["cat_done_state"] = {}
        for c, is_done in cat_done.items():
            prev = st.session_state["cat_done_state"].get(c, False)
            if is_done and not prev:
                st.session_state["cat_done_state"][c] = True
                try:
                    send_email(cfg, f"S-STEM Prep: Category Complete — {c}", f"<p>All tasks in <b>{c}</b> are complete.</p>")
                except Exception as e:
                    st.warning(f"Category alert failed: {e}")
            elif not is_done:
                st.session_state["cat_done_state"][c] = False

        if (base['Status']=="Complete").all() and not st.session_state.get("ALL_DONE", False):
            st.session_state["ALL_DONE"] = True
            try:
                send_email(cfg, "S-STEM Prep: ALL TASKS COMPLETE 🎉", "<p>All preparatory tasks are complete. Ready to submit.</p>")
            except Exception as e:
                st.warning(f"All-done alert failed: {e}")
        elif not (base['Status']=="Complete").all():
            st.session_state["ALL_DONE"] = False

        st.experimental_rerun()

with tab_settings:
    cfg = load_config()
    st.markdown("#### Notifications")
    recip = st.text_area("Recipients (comma-separated)", value=",".join(cfg["recipients"]))
    c1,c2,c3,c4 = st.columns(4)
    with c1:
        host = st.text_input("SMTP Host", value=cfg["smtp"]["host"])
    with c2:
        port = st.number_input("SMTP Port", value=cfg["smtp"]["port"])
    with c3:
        username = st.text_input("SMTP Username", value=cfg["smtp"]["username"])
    with c4:
        password = st.text_input("SMTP Password", value=cfg["smtp"]["password"], type="password")
    if st.button("Save Notification Settings"):
        cfg["recipients"] = [x.strip() for x in recip.split(",") if x.strip()]
        cfg["smtp"].update({"host":host,"port":int(port),"username":username,"password":password})
        save_config(cfg)
        st.success("Saved notification settings.")

    st.markdown("---")
    st.markdown("#### Mentoring Handbook Generator")
    col1, col2 = st.columns(2)
    with col1:
        mh_meet = st.text_input("Meeting cadence", "Biweekly, 45 minutes")
        mh_comm = st.text_input("Communication norms", "Email within 2 business days; Teams for quick updates")
        mh_track = st.text_input("Progress tracking", "Shared notes; action items; quarterly rubric")
    with col2:
        mh_incl = st.text_input("Equity & inclusion", "Culturally responsive mentoring; accessibility accommodations")
        mh_risk = st.text_input("Risk & reporting", "COI disclosure; escalation path PI → EAB → Dean")
        mh_sign = st.text_input("Signature lines", "Mentor / Scholar / PI")
    if st.button("Generate Mentoring Handbook (DOCX)"):
        doc = Document()
        doc.add_heading("S-STEM Mentoring Handbook", level=1)
        doc.add_paragraph("Purpose: Define expectations and structures for mentoring S-STEM scholars, including faculty mentors and industry supervisors.")
        doc.add_heading("Roles & Expectations", level=2)
        doc.add_paragraph("- Faculty Mentor: Academic guidance, career planning, research exposure.")
        doc.add_paragraph("- Industry Supervisor: Workplace mentorship, project scoping, professional skills.")
        doc.add_paragraph("- Scholar: Active engagement, meeting preparation, progress documentation.")
        doc.add_heading("Meeting Cadence", level=2); doc.add_paragraph(mh_meet)
        doc.add_heading("Communication Norms", level=2); doc.add_paragraph(mh_comm)
        doc.add_heading("Progress Tracking", level=2); doc.add_paragraph(mh_track)
        doc.add_heading("Equity & Inclusion", level=2); doc.add_paragraph(mh_incl)
        doc.add_heading("Risk Management & Reporting", level=2); doc.add_paragraph(mh_risk)
        doc.add_heading("Signatures", level=2); doc.add_paragraph(mh_sign)
        out = os.path.join(BASE, "Mentoring_Handbook.docx")
        doc.save(out)
        with open(out, "rb") as f:
            st.download_button("Download Mentoring_Handbook.docx", data=f.read(), file_name="Mentoring_Handbook.docx")
